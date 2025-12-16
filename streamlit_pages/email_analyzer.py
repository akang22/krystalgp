"""Streamlit page for analyzing individual emails with all parsers."""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import io
import os
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from typing import Optional

import pandas as pd
import streamlit as st
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image

# Load secrets from Streamlit secrets.toml into environment
# This allows parsers to work with st.secrets or .env files
if hasattr(st, "secrets"):
    for key in st.secrets.keys():
        if key not in os.environ:
            os.environ[key] = st.secrets[key]

from email_parser.ensemble_parser import EnsembleParser
from email_parser.layout_attachment_parser import LayoutLLMParser
from email_parser.llm_body_parser import LLMBodyParser
from email_parser.ocr_attachment_parser import OCRAttachmentParser

# Paths
WORKSPACE = Path(__file__).parent.parent
SAMPLE_EMAILS_DIR = WORKSPACE / "sample_emails"
RESULTS_CSV = WORKSPACE / "results.csv"


@st.cache_resource
def get_parsers():
    """Initialize all parsers (cached)."""
    parsers = {}

    try:
        parsers["LLM Body"] = LLMBodyParser()
    except Exception as e:
        st.warning("LLM parser not available (set OPENAI_API_KEY)")

    try:
        parsers["OCR + LLM"] = OCRAttachmentParser()
    except ValueError as e:
        # API key error
        st.warning("OCR + LLM parser not available (set OPENAI_API_KEY)")
    except RuntimeError as e:
        # Tesseract error
        error_msg = str(e)
        if "Tesseract" in error_msg:
            st.error(f"⚠️ OCR + LLM parser not available: {error_msg}")
        else:
            st.warning(f"OCR + LLM parser not available: {error_msg}")
    except Exception as e:
        # Other errors
        st.warning(f"OCR + LLM parser not available: {str(e)[:200]}")

    try:
        parsers["Layout Vision"] = LayoutLLMParser()
    except Exception as e:
        st.warning("Layout Vision parser not available (set OPENAI_API_KEY)")

    try:
        parsers["Final Results"] = EnsembleParser(
            use_llm=True, use_vision=True, use_ocr=False, results_csv_path=RESULTS_CSV
        )
    except Exception as e:
        st.warning(f"Final Results parser not available: {e}")

    return parsers


def display_email_metadata(email_data):
    """Display email metadata."""
    st.subheader("📧 Email Metadata")

    col1, col2 = st.columns(2)

    with col1:
        st.write("**From:**", email_data.sender or "N/A")
        st.write("**Date:**", str(email_data.date) if email_data.date else "N/A")

    with col2:
        st.write("**To:**", ", ".join(email_data.recipients[:2]) or "N/A")
        if len(email_data.recipients) > 2:
            st.write(f"*... and {len(email_data.recipients) - 2} more*")

    st.write("**Subject:**", email_data.subject or "N/A")


def display_attachments(email_data):
    """Display attachment information."""
    st.subheader("📎 Attachments")

    if not email_data.attachments:
        st.info("No attachments found")
        return

    st.write(f"**Total:** {len(email_data.attachments)} file(s)")

    attachment_data = []
    for att in email_data.attachments:
        attachment_data.append(
            {
                "Filename": att.filename,
                "Type": att.content_type or "Unknown",
                "Size": f"{att.size_bytes / 1024:.1f} KB",
            }
        )

    st.dataframe(pd.DataFrame(attachment_data), width="stretch")


def display_email_body(email_data):
    """Display email body."""
    st.subheader("📄 Email Body")

    body = email_data.body_plain or email_data.body_html or ""

    if not body:
        st.info("No body text found")
        return

    # Show first 2000 chars in expander
    with st.expander("View Email Content", expanded=False):
        st.text_area("Body", body, height=300)


def display_parser_results(results):
    """Display results from all parsers."""
    st.subheader("🔍 Parser Results Breakdown")

    # Create comparison table
    comparison_data = []
    final_results_data = None

    # Separate Final Results to show at the bottom
    for parser_name, result in results.items():
        if parser_name == "Final Results":
            # Store Final Results separately to add at the end
            if result:
                opp = result.opportunity

                # Format EBITDA with options count
                ebitda_str = f"${opp.ebitda_millions:.2f}M" if opp.ebitda_millions else "Not found"
                if hasattr(opp, "ebitda_options") and opp.ebitda_options:
                    ebitda_str += f" ({len(opp.ebitda_options)} options)"

                # Format location with options count
                location_str = opp.hq_location or "Not found"
                if hasattr(opp, "location_options") and opp.location_options:
                    location_str += f" ({len(opp.location_options)} options)"

                # Format company with options count
                company_str = opp.company_name or "Not found"
                if hasattr(opp, "company_options") and opp.company_options:
                    company_str += f" ({len(opp.company_options)} options)"

                # Format sector with options count
                sector_str = opp.sector or "Not found"
                if hasattr(opp, "sector_options") and opp.sector_options:
                    sector_str += f" ({len(opp.sector_options)} options)"

                final_results_data = {
                    "Parser": parser_name,
                    "EBITDA": ebitda_str,
                    "Company": company_str,
                    "HQ Location": location_str,
                    "Sector": sector_str,
                    "Source": result.extraction_source,
                    "Time (s)": f"{result.processing_time_seconds:.2f}",
                }
            else:
                final_results_data = {
                    "Parser": "Final Results",
                    "EBITDA": "Error",
                    "Company": "Error",
                    "HQ Location": "Error",
                    "Sector": "Error",
                    "Source": "N/A",
                    "Time (s)": "N/A",
                }
            continue  # Skip adding to main list, will add at end

        # Process all other parsers
        if result:
            opp = result.opportunity

            # Format EBITDA with options count
            ebitda_str = f"${opp.ebitda_millions:.2f}M" if opp.ebitda_millions else "Not found"
            if hasattr(opp, "ebitda_options") and opp.ebitda_options:
                ebitda_str += f" ({len(opp.ebitda_options)} options)"

            # Format location with options count
            location_str = opp.hq_location or "Not found"
            if hasattr(opp, "location_options") and opp.location_options:
                location_str += f" ({len(opp.location_options)} options)"

            # Format company with options count
            company_str = opp.company_name or "Not found"
            if hasattr(opp, "company_options") and opp.company_options:
                company_str += f" ({len(opp.company_options)} options)"

            # Format sector with options count
            sector_str = opp.sector or "Not found"
            if hasattr(opp, "sector_options") and opp.sector_options:
                sector_str += f" ({len(opp.sector_options)} options)"

            comparison_data.append(
                {
                    "Parser": parser_name,
                    "EBITDA": ebitda_str,
                    "Company": company_str,
                    "HQ Location": location_str,
                    "Sector": sector_str,
                    "Source": result.extraction_source,
                    "Time (s)": f"{result.processing_time_seconds:.2f}",
                }
            )
        else:
            comparison_data.append(
                {
                    "Parser": parser_name,
                    "EBITDA": "Error",
                    "Company": "Error",
                    "HQ Location": "Error",
                    "Sector": "Error",
                    "Source": "N/A",
                    "Time (s)": "N/A",
                }
            )

    # Add Final Results at the bottom if it exists
    if final_results_data:
        comparison_data.append(final_results_data)

    df = pd.DataFrame(comparison_data)
    st.dataframe(df, width="stretch", hide_index=True)


def display_confidence_calculation(results):
    """Display ensemble selection logic (NOT averaging)."""
    st.subheader("🎯 Ensemble Selection Logic")

    st.info("**Note:** The ensemble SELECTS the best value, it does NOT average them!")

    # Filter parsers with valid EBITDA
    valid_results = []
    for parser_name, result in results.items():
        if result and result.opportunity.ebitda_millions and parser_name != "Ensemble (Confidence)":
            valid_results.append((parser_name, result))

    if not valid_results:
        st.warning("No valid EBITDA values found")
        return

    # Define weights
    parser_weights = {
        "LLM Body": 1.0,
        "OCR + LLM": 0.5,
        "Layout Vision": 0.9,
    }

    source_weights = {
        "body": 1.0,
        "attachment": 1.2,
        "both": 1.1,
    }

    raw_text_bonus = 1.1

    # Build calculation table
    calc_data = []
    total_weighted = 0
    total_weight = 0

    for parser_name, result in valid_results:
        opp = result.opportunity
        ebitda = opp.ebitda_millions

        parser_weight = parser_weights.get(parser_name, 0.5)
        source_weight = source_weights.get(result.extraction_source, 1.0)
        has_raw = bool(opp.raw_ebitda_text)

        final_weight = parser_weight * source_weight * (raw_text_bonus if has_raw else 1.0)
        weighted_value = ebitda * final_weight

        calc_data.append(
            {
                "Parser": parser_name,
                "EBITDA": f"${ebitda:.2f}M",
                "Base Weight": parser_weight,
                "Source": result.extraction_source,
                "Source Mult": f"{source_weight}×",
                "Raw Text Bonus": "✓" if has_raw else "✗",
                "Final Weight": f"{final_weight:.3f}",
                "Weighted Value": f"{weighted_value:.3f}",
            }
        )

        total_weighted += weighted_value
        total_weight += final_weight

    # Display calculation table
    st.dataframe(pd.DataFrame(calc_data), width="stretch", hide_index=True)

    # Show selection logic
    st.markdown("### 🎯 Selection Logic")

    # Check for fuzzy consensus
    ebitda_values = [result.opportunity.ebitda_millions for _, result in valid_results]

    # Count occurrences (for fuzzy matching)
    from collections import Counter

    value_counts = Counter(ebitda_values)
    most_common = value_counts.most_common(1)[0] if value_counts else (None, 0)

    if most_common[1] >= 2:
        st.success(
            f"""
        ✅ **Fuzzy Consensus Found!**
        
        **${most_common[0]:.2f}M** appears {most_common[1]} times (majority)
        
        → **SELECTED: ${most_common[0]:.2f}M**
        """
        )
    else:
        # Find highest confidence
        best_parser = max(
            valid_results, key=lambda x: calc_data[valid_results.index(x)]["Final Weight"]
        )
        best_ebitda = best_parser[1].opportunity.ebitda_millions
        best_name = best_parser[0]

        st.warning(
            f"""
        ⚠️ **No Consensus - Using Confidence Selection**
        
        Highest confidence: **{best_name}**
        
        → **SELECTED: ${best_ebitda:.2f}M**
        """
        )

    # Show what ensemble returned
    ensemble_result = results.get("Ensemble (Confidence)")
    if ensemble_result and ensemble_result.opportunity.ebitda_millions:
        final_value = ensemble_result.opportunity.ebitda_millions
        method = ensemble_result.opportunity.raw_ebitda_text or "Unknown"

        col1, col2 = st.columns(2)
        with col1:
            st.metric("🎯 Ensemble Selected", f"${final_value:.2f}M")
        with col2:
            st.metric("Selection Method", method.replace("[", "").replace("]", ""))


def display_detailed_results(results):
    """Display detailed results for each parser."""
    st.subheader("📋 Detailed Parser Results")

    for parser_name, result in results.items():
        with st.expander(f"**{parser_name}**"):
            if not result:
                st.error("Parser failed")
                continue

            opp = result.opportunity

            col1, col2 = st.columns(2)

            with col1:
                st.write(
                    "**EBITDA:**",
                    f"${opp.ebitda_millions:.2f}M" if opp.ebitda_millions else "Not found",
                )
                st.write("**Company:**", opp.company_name or "Not found")
                st.write("**HQ Location:**", opp.hq_location or "Not found")
                st.write("**Sector:**", opp.sector or "Not found")

            with col2:
                st.write("**Source Domain:**", opp.source_domain or "Not found")
                st.write("**Recipient:**", opp.recipient or "Not found")
                st.write("**Processing Time:**", f"{result.processing_time_seconds:.2f}s")
                st.write("**Extraction Source:**", result.extraction_source)

            # Show multiple options with confidence scores
            if hasattr(opp, "ebitda_options") and opp.ebitda_options:
                st.markdown("**💡 EBITDA Options (All Candidates):**")
                ebitda_df = []
                for opt in sorted(opp.ebitda_options, key=lambda x: x.confidence, reverse=True):
                    ebitda_df.append(
                        {
                            "Value": f"${opt.value}M",
                            "Confidence": f"{opt.confidence:.0%}",
                            "Source": opt.source,
                            "Raw Text": opt.raw_text or "N/A",
                        }
                    )
                if ebitda_df:
                    st.dataframe(ebitda_df, width="stretch", hide_index=True)

            if hasattr(opp, "location_options") and opp.location_options:
                st.markdown("**💡 Location Options (All Candidates):**")
                loc_df = []
                for opt in sorted(opp.location_options, key=lambda x: x.confidence, reverse=True):
                    loc_df.append(
                        {
                            "Value": opt.value,
                            "Confidence": f"{opt.confidence:.0%}",
                            "Source": opt.source,
                            "Raw Text": opt.raw_text or "N/A",
                        }
                    )
                if loc_df:
                    st.dataframe(loc_df, width="stretch", hide_index=True)

            if hasattr(opp, "company_options") and opp.company_options:
                st.markdown("**💡 Company Options (All Candidates):**")
                comp_df = []
                for opt in sorted(opp.company_options, key=lambda x: x.confidence, reverse=True):
                    comp_df.append(
                        {
                            "Value": opt.value,
                            "Confidence": f"{opt.confidence:.0%}",
                            "Source": opt.source,
                            "Raw Text": opt.raw_text or "N/A",
                        }
                    )
                if comp_df:
                    st.dataframe(comp_df, width="stretch", hide_index=True)

            if hasattr(opp, "sector_options") and opp.sector_options:
                st.markdown("**💡 Sector Options (All Candidates):**")
                sector_df = []
                for opt in sorted(opp.sector_options, key=lambda x: x.confidence, reverse=True):
                    sector_df.append(
                        {
                            "Value": opt.value,
                            "Confidence": f"{opt.confidence:.0%}",
                            "Source": opt.source,
                            "Raw Text": opt.raw_text or "N/A",
                        }
                    )
                if sector_df:
                    st.dataframe(sector_df, width="stretch", hide_index=True)

            if opp.raw_ebitda_text and not opp.ebitda_options:
                st.write("**Raw EBITDA Text:**")
                st.code(opp.raw_ebitda_text)


def display_pdf_attachment(attachment):
    """Display PDF attachment."""
    try:
        # Convert PDF to images
        images = convert_from_bytes(attachment.content, dpi=150)

        st.markdown(f"**{attachment.filename}** ({attachment.size_bytes / 1024:.1f} KB)")

        # Display first 3 pages
        for i, img in enumerate(images[:3]):
            st.image(img, caption=f"Page {i+1}", use_container_width=True)

        if len(images) > 3:
            st.info(f"Showing first 3 pages of {len(images)} total pages")

    except Exception as e:
        st.error(f"Failed to display PDF: {e}")


def display_image_attachment(attachment):
    """Display image attachment."""
    try:
        image = Image.open(io.BytesIO(attachment.content))
        st.markdown(f"**{attachment.filename}** ({attachment.size_bytes / 1024:.1f} KB)")
        st.image(image, use_container_width=True)
    except Exception as e:
        st.error(f"Failed to display image: {e}")


def display_docx_attachment(attachment):
    """Display Word document as text."""
    try:
        doc = Document(io.BytesIO(attachment.content))
        st.markdown(f"**{attachment.filename}** ({attachment.size_bytes / 1024:.1f} KB)")

        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)

        text_content = "\n\n".join(full_text)

        if text_content:
            st.text_area("Document Content", text_content, height=400, label_visibility="collapsed")
        else:
            st.warning("No text content found in document")

    except Exception as e:
        st.error(f"Failed to display Word document: {e}")


def display_text_attachment(attachment):
    """Display text-based attachments."""
    try:
        # Try to decode as text
        text_content = attachment.content.decode("utf-8", errors="ignore")

        st.markdown(f"**{attachment.filename}** ({attachment.size_bytes / 1024:.1f} KB)")
        st.text_area(
            "File Content",
            text_content[:5000],  # First 5000 chars
            height=300,
            label_visibility="collapsed",
        )

        if len(text_content) > 5000:
            st.info(f"Showing first 5000 characters of {len(text_content)} total")

    except Exception as e:
        st.error(f"Failed to display text file: {e}")


def display_attachments_visual(email_data):
    """Display attachments with preview."""
    st.subheader("📎 Attachments")

    if not email_data.attachments:
        st.info("No attachments found")
        return

    st.write(f"**Total:** {len(email_data.attachments)} file(s)")

    # Display each attachment
    for att in email_data.attachments:
        filename_lower = att.filename.lower()

        with st.expander(f"📄 {att.filename} ({att.size_bytes / 1024:.1f} KB)"):
            if filename_lower.endswith(".pdf"):
                display_pdf_attachment(att)
            elif any(
                filename_lower.endswith(ext)
                for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"]
            ):
                display_image_attachment(att)
            elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
                display_docx_attachment(att)
            elif any(filename_lower.endswith(ext) for ext in [".txt", ".csv", ".log", ".md"]):
                display_text_attachment(att)
            else:
                # Try to display as text
                st.info(f"Type: {att.content_type or 'Unknown'}")
                try:
                    display_text_attachment(att)
                except:
                    st.warning("Preview not available for this file type")


def calculate_investment_criteria_fit(
    hq_location: Optional[str], ebitda_millions: Optional[float]
) -> str:
    """Calculate if investment opportunity fits criteria.

    Criteria: HQ = Western Canada (BC or Alberta) && EBITDA between 2-8M

    Args:
        hq_location: Headquarters location
        ebitda_millions: EBITDA in millions

    Returns:
        "Yes" if criteria met, "No" otherwise
    """
    if not hq_location or ebitda_millions is None:
        return "No"

    # Hardcoded list of 25 populous cities in BC and Alberta
    # BC cities (15): Vancouver, Victoria, Surrey, Burnaby, Richmond, Coquitlam, 
    # Langley, Abbotsford, North Vancouver, West Vancouver, Kelowna, Kamloops, 
    # Nanaimo, Prince George, Chilliwack
    # Alberta cities (10): Calgary, Edmonton, Red Deer, Lethbridge, St. Albert, 
    # Medicine Hat, Grande Prairie, Airdrie, Spruce Grove, Fort McMurray
    western_canada_cities = [
        # BC cities
        "VANCOUVER", "VICTORIA", "SURREY", "BURNABY", "RICHMOND",
        "COQUITLAM", "LANGLEY", "ABBOTSFORD", "NORTH VANCOUVER", "WEST VANCOUVER",
        "KELOWNA", "KAMLOOPS", "NANAIMO", "PRINCE GEORGE", "CHILLIWACK",
        # Alberta cities
        "CALGARY", "EDMONTON", "RED DEER", "LETHBRIDGE", "ST. ALBERT",
        "MEDICINE HAT", "GRANDE PRAIRIE", "AIRDRIE", "SPRUCE GROVE", "FORT MCMURRAY",
    ]

    # Check if location is in Western Canada (BC or Alberta)
    location_upper = hq_location.upper()
    
    # Check for province abbreviations/names
    is_western_canada = (
        "BC" in location_upper
        or "BRITISH COLUMBIA" in location_upper
        or "ALBERTA" in location_upper
        or "AB" in location_upper
    )
    
    # Check for regional terms indicating Western Canada
    if not is_western_canada:
        western_terms = [
            "WEST COAST",
            "WESTERN CANADA",
            "WESTERN CANADIAN",
            "BC-BASED",
            "ALBERTA-BASED",
            "PACIFIC COAST",  # BC is on the Pacific
        ]
        for term in western_terms:
            if term in location_upper:
                is_western_canada = True
                break
    
    # Check for major cities (case-insensitive)
    if not is_western_canada:
        for city in western_canada_cities:
            if city in location_upper:
                is_western_canada = True
                break

    # Check if EBITDA is between 2-8M
    ebitda_in_range = 2.0 <= ebitda_millions <= 8.0

    return "Yes" if (is_western_canada and ebitda_in_range) else "No"


def display_summary_table(email_data, results: dict):
    """Display summary table at the top with key investment opportunity data.

    Args:
        email_data: EmailData object
        results: Dictionary of parser results
    """
    # Get Final Results if available, otherwise use first available result
    final_result = results.get("Final Results")
    llm_body_result = results.get("LLM Body")  # Prefer LLM Body for description

    if not final_result:
        # Try to get any result
        for parser_name, result in results.items():
            if result and result.opportunity:
                final_result = result
                break

    if not final_result or not final_result.opportunity:
        st.warning("No parser results available for summary table")
        return

    opp = final_result.opportunity

    # Get description from LLM Body parser if available (it generates the description)
    if llm_body_result and llm_body_result.opportunity:
        llm_opp = llm_body_result.opportunity
        if hasattr(llm_opp, "description") and llm_opp.description:
            opp.description = llm_opp.description  # Override with LLM-generated description

    # Extract data
    date_received = email_data.date.strftime("%Y-%m-%d") if email_data.date else "N/A"
    company_name = opp.company_name or "N/A"
    sector = opp.sector or "N/A"
    # Use LLM-generated description, fallback to subject (with defensive check for old cached data)
    description = getattr(opp, "description", None) or email_data.subject or "N/A"
    ebitda = f"${opp.ebitda_millions:.2f}M" if opp.ebitda_millions is not None else "N/A"
    hq_location = opp.hq_location or "N/A"
    source = opp.source_domain or "N/A"

    # Extract receiver username from recipient email
    receiver = "N/A"
    recipient_email = opp.recipient or (email_data.recipients[0] if email_data.recipients else None)
    if recipient_email:
        # First extract the actual email address (in case it's in format "Name <email@domain.com>")
        import email.utils
        parsed_addr = email.utils.parseaddr(recipient_email)
        email_addr = parsed_addr[1] if parsed_addr[1] else recipient_email
        
        # Then extract username part (before @)
        if "@" in email_addr:
            receiver = email_addr.split("@")[0]
        else:
            receiver = email_addr

    # Calculate investment criteria fit
    investment_fit = calculate_investment_criteria_fit(opp.hq_location, opp.ebitda_millions)

    # Create table data
    table_data = {
        "Date Received": [date_received],
        "Company / Project Name": [company_name],
        "Sector": [sector],
        "Description": [description],
        "LTM EBITDA ($M)": [ebitda],
        "HQ Location": [hq_location],
        "Source": [source],
        "Receiver": [receiver],
        "Investment Criteria Fit?": [investment_fit],
    }

    df = pd.DataFrame(table_data)

    # Display table with styling
    st.subheader("📊 Investment Opportunity Summary")

    # Style the dataframe with color coding for Investment Criteria Fit column
    def style_status(val):
        if val == "Yes":
            return "background-color: #90EE90; color: #000000"  # Light green
        else:
            return "background-color: #FFB6C1; color: #000000"  # Light red

    # Apply styling to Investment Criteria Fit column
    styled_df = df.style.applymap(style_status, subset=["Investment Criteria Fit?"])

    st.dataframe(styled_df, width="stretch", hide_index=True, use_container_width=True)

    st.divider()


def main():
    """Main Streamlit app."""
    # Get list of emails
    # Email input method selection
    st.subheader("📨 Select Email to Analyze")
    
    input_method = st.radio(
        "Choose input method:",
        ["Sample Email", "Upload .eml File"],
        horizontal=True,
    )

    email_path = None
    selected_email = None
    uploaded_file = None

    if input_method == "Sample Email":
        email_files = sorted([f.name for f in SAMPLE_EMAILS_DIR.glob("*.msg")])

        if not email_files:
            st.error(f"No .msg files found in {SAMPLE_EMAILS_DIR}")
            return

        selected_email = st.selectbox(
            "Choose an email:",
            email_files,
            index=(
                email_files.index(
                    "FW Project Gravy - Franchise QSR Portfolio Acquisition Opportunity.msg"
                )
                if "FW Project Gravy - Franchise QSR Portfolio Acquisition Opportunity.msg"
                in email_files
                else 0
            ),
            label_visibility="collapsed",
        )

        email_path = SAMPLE_EMAILS_DIR / selected_email
    else:
        # File upload
        uploaded_file = st.file_uploader(
            "Upload .eml file",
            type=["eml"],
            help="Upload an .eml email file for analysis",
        )

        if uploaded_file is not None:
            # Save uploaded file to a temporary location
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".eml") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                email_path = Path(tmp_file.name)
                selected_email = uploaded_file.name

            st.success(f"✅ Uploaded: {uploaded_file.name}")
        else:
            st.info("Please upload an .eml file to analyze")
            return

    if not email_path or not email_path.exists():
        st.error("Email file not found")
        return

    st.divider()

    # Initialize parsers
    parsers = get_parsers()

    if not parsers:
        st.error("No parsers available. Check your configuration.")
        return

    # Initialize session state for caching
    if "cached_results" not in st.session_state:
        st.session_state.cached_results = {}
    if "cached_email_data" not in st.session_state:
        st.session_state.cached_email_data = {}

    # Check if we have cached results for this email
    # For uploaded files, use a unique cache key based on file content hash
    if uploaded_file is not None:
        import hashlib
        file_hash = hashlib.md5(uploaded_file.getvalue()).hexdigest()[:8]
        cache_key = f"uploaded_{file_hash}_{selected_email}"
    else:
        cache_key = selected_email
    
    use_cached = (
        cache_key in st.session_state.cached_results
        and cache_key in st.session_state.cached_email_data
        and input_method == "Sample Email"  # Don't cache uploaded files
    )

    if use_cached:
        # Load from cache
        st.info("📦 Loaded results from cache (click 'Reparse' to refresh)")
        email_data = st.session_state.cached_email_data[cache_key]
        results = st.session_state.cached_results[cache_key]
    else:
        # Parse email
        with st.spinner("Parsing email..."):
            # Get email metadata first
            try:
                first_parser = list(parsers.values())[0]
                # Detect file type and use appropriate extraction method
                if email_path.suffix.lower() == ".eml":
                    email_data = first_parser.extract_eml_file(email_path)
                else:
                    email_data = first_parser.extract_msg_file(email_path)
            except Exception as e:
                st.error(f"Failed to read email: {e}")
                import traceback
                st.code(traceback.format_exc())
                return

            # Run all parsers in parallel
            results = {}

            progress_bar = st.progress(0)
            status_text = st.empty()

            error_log = []

            # Run parsers in parallel using ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=len(parsers)) as executor:
                # Submit all parser tasks
                future_to_parser = {
                    executor.submit(parser.parse, email_path): parser_name
                    for parser_name, parser in parsers.items()
                }

                # Collect results as they complete
                completed = 0
                for future in as_completed(future_to_parser):
                    parser_name = future_to_parser[future]
                    completed += 1
                    status_text.text(f"Running {parser_name}... ({completed}/{len(parsers)})")

                    try:
                        result = future.result()
                        results[parser_name] = result

                        # Debug: show what was extracted
                        if result and result.opportunity:
                            opp = result.opportunity
                            ebitda_str = (
                                f"${opp.ebitda_millions:.2f}M" if opp.ebitda_millions else "None"
                            )
                            status_text.text(
                                f"✓ {parser_name}: EBITDA={ebitda_str}, Company={opp.company_name or 'None'}"
                            )

                            # Log errors if any
                            if result.errors:
                                error_log.append(f"{parser_name}: {', '.join(result.errors)}")

                    except Exception as e:
                        error_msg = f"{parser_name} failed: {str(e)}"
                        error_log.append(error_msg)
                        st.error(f"❌ {error_msg}")
                        results[parser_name] = None

                    progress_bar.progress(completed / len(parsers))

            status_text.text("✅ Parsing complete!")

            # Show errors if any
            if error_log:
                with st.expander("⚠️ Errors/Warnings"):
                    for err in error_log:
                        st.warning(err)

            # Cache results
            st.session_state.cached_results[cache_key] = results
            st.session_state.cached_email_data[cache_key] = email_data

    # Add reparse button (only show for sample emails, not uploaded files)
    if input_method == "Sample Email":
        if st.button("🔄 Reparse Email", help="Clear cache and reparse this email"):
            # Clear all cache entries for this email to force fresh parsing
            if cache_key in st.session_state.cached_results:
                del st.session_state.cached_results[cache_key]
            if cache_key in st.session_state.cached_email_data:
                del st.session_state.cached_email_data[cache_key]
            # Rerun will trigger fresh parsing since cache is now empty
            st.rerun()

    # Display summary table at the very top
    display_summary_table(email_data, results)

    # Display email info
    display_email_metadata(email_data)

    st.divider()

    # Email body
    display_email_body(email_data)

    st.divider()

    # Attachments with visual display
    display_attachments_visual(email_data)

    st.divider()

    # Display results
    display_parser_results(results)

    st.divider()

    # Detailed results
    display_detailed_results(results)


if __name__ == "__main__":
    main()
